"""
Document Exporter for multiple file formats
Supports PDF, DOCX, HTML, and Markdown output
"""

import os
import io
import base64
from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime
import tempfile

class DocumentExporter:
    """Handles export of documents to various formats"""
    
    def __init__(self, company_config: Dict[str, Any]):
        self.company_config = company_config
        self.supported_formats = ['markdown', 'html', 'pdf', 'docx']
    
    def export_document(self, content: str, project_data: Dict[str, Any], 
                       output_format: str, output_path: str) -> bool:
        """Export document to specified format"""
        try:
            if output_format.lower() not in self.supported_formats:
                raise ValueError(f"Unsupported format: {output_format}")
            
            # Process content with company branding
            processed_content = self._add_company_branding(content, project_data)
            
            if output_format.lower() == 'markdown':
                return self._export_markdown(processed_content, output_path)
            elif output_format.lower() == 'html':
                return self._export_html(processed_content, output_path)
            elif output_format.lower() == 'pdf':
                return self._export_pdf(processed_content, output_path)
            elif output_format.lower() == 'docx':
                return self._export_docx(processed_content, output_path)
            
        except Exception as e:
            print(f"Export error: {e}")
            return False
    
    def _add_company_branding(self, content: str, project_data: Dict[str, Any]) -> str:
        """Add company branding to the document"""
        company = self.company_config
        
        # Create header with company information
        header = self._create_document_header(project_data)
        
        # Create footer with company information
        footer = self._create_document_footer()
        
        # Combine header, content, and footer
        branded_content = f"{header}\n\n{content}\n\n{footer}"
        
        return branded_content
    
    def _create_document_header(self, project_data: Dict[str, Any]) -> str:
        """Create document header with company branding"""
        company = self.company_config
        
        header = f"""# {project_data.get('project_name', 'Project Proposal')}

---

**{company.get('name', 'Company Name')}**  
*{company.get('tagline', 'Your trusted partner')}*

**Industry:** {company.get('industry', 'Technology Solutions')}  
**Website:** {company.get('website', '')}  
**Email:** {company.get('email', '')}  
**Phone:** {company.get('phone', '')}

---

**Prepared for:** {project_data.get('client_name', 'Valued Client')}  
**Date:** {datetime.now().strftime('%B %d, %Y')}  
**Project Type:** {project_data.get('project_type', 'Software')}  
**Document Version:** {project_data.get('version', '1.0')}  
**Prepared by:** {project_data.get('author', self.company_config.get('default_author', 'Project Team'))}

---
"""
        return header
    
    def _create_document_footer(self) -> str:
        """Create document footer with company information"""
        company = self.company_config
        
        footer = f"""---

## About {company.get('name', 'Our Company')}

{company.get('name', 'Our Company')} is a leading {company.get('industry', 'technology')} company specializing in:

"""
        
        # Add specializations
        specializations = company.get('specializations', [])
        for spec in specializations:
            footer += f"- {spec}\n"
        
        footer += f"""
**Contact Information:**

{company.get('address', '')}

**Phone:** {company.get('phone', '')}  
**Email:** {company.get('email', '')}  
**Website:** {company.get('website', '')}

**Legal Information:**  
Registration Number: {company.get('registration_number', '')}  
Tax ID: {company.get('tax_id', '')}

---

*This document is confidential and proprietary to {company.get('name', 'Our Company')}. Any reproduction or distribution without written permission is strictly prohibited.*

**Document Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
"""
        return footer
    
    def _export_markdown(self, content: str, output_path: str) -> bool:
        """Export as Markdown"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return True
        except Exception as e:
            print(f"Markdown export error: {e}")
            return False
    
    def _export_html(self, content: str, output_path: str) -> bool:
        """Export as HTML with styling"""
        try:
            # Convert markdown to HTML (basic conversion)
            html_content = self._markdown_to_html(content)
            
            # Add CSS styling
            styled_html = self._add_html_styling(html_content)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(styled_html)
            return True
        except Exception as e:
            print(f"HTML export error: {e}")
            return False
    
    def _export_pdf(self, content: str, output_path: str) -> bool:
        """Export as PDF using reportlab"""
        try:
            # Try to import reportlab
            try:
                from reportlab.lib.pagesizes import letter, A4
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                from reportlab.lib import colors
            except ImportError:
                print("PDF export requires reportlab: pip install reportlab")
                return False
            
            # Create PDF document
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                textColor=colors.darkblue
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                spaceAfter=12,
                spaceBefore=12,
                textColor=colors.darkblue
            )
            
            # Convert markdown content to PDF elements
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 12))
                elif line.startswith('# '):
                    story.append(Paragraph(line[2:], title_style))
                elif line.startswith('## '):
                    story.append(Paragraph(line[3:], heading_style))
                elif line.startswith('### '):
                    story.append(Paragraph(line[4:], styles['Heading3']))
                elif line.startswith('**') and line.endswith('**'):
                    story.append(Paragraph(f"<b>{line[2:-2]}</b>", styles['Normal']))
                elif line.startswith('*') and line.endswith('*'):
                    story.append(Paragraph(f"<i>{line[1:-1]}</i>", styles['Normal']))
                elif line.startswith('- '):
                    story.append(Paragraph(f"• {line[2:]}", styles['Normal']))
                elif line.startswith('---'):
                    story.append(Spacer(1, 20))
                else:
                    if line:
                        story.append(Paragraph(line, styles['Normal']))
            
            # Build PDF
            doc.build(story)
            return True
            
        except Exception as e:
            print(f"PDF export error: {e}")
            return False
    
    def _export_docx(self, content: str, output_path: str) -> bool:
        """Export as DOCX using python-docx"""
        try:
            # Try to import python-docx
            try:
                from docx import Document
                from docx.shared import Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.shared import RGBColor
            except ImportError:
                print("DOCX export requires python-docx: pip install python-docx")
                return False
            
            # Create document
            doc = Document()
            
            # Add company logo if available
            logo_path = self.company_config.get('logo_path', '')
            if logo_path and os.path.exists(logo_path):
                try:
                    doc.add_picture(logo_path, width=Inches(2))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except:
                    pass  # Skip logo if there's an issue
            
            # Process content
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                elif line.startswith('# '):
                    heading = doc.add_heading(line[2:], 1)
                    heading.runs[0].font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
                elif line.startswith('## '):
                    heading = doc.add_heading(line[3:], 2)
                    heading.runs[0].font.color.rgb = RGBColor(0, 0, 139)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], 3)
                elif line.startswith('**') and line.endswith('**'):
                    p = doc.add_paragraph()
                    run = p.add_run(line[2:-2])
                    run.bold = True
                elif line.startswith('*') and line.endswith('*'):
                    p = doc.add_paragraph()
                    run = p.add_run(line[1:-1])
                    run.italic = True
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('---'):
                    doc.add_page_break()
                else:
                    if line:
                        doc.add_paragraph(line)
            
            # Save document
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"DOCX export error: {e}")
            return False
    
    def _markdown_to_html(self, markdown_content: str) -> str:
        """Basic markdown to HTML conversion"""
        lines = markdown_content.split('\n')
        html_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                html_lines.append('<br>')
            elif line.startswith('# '):
                html_lines.append(f'<h1>{line[2:]}</h1>')
            elif line.startswith('## '):
                html_lines.append(f'<h2>{line[3:]}</h2>')
            elif line.startswith('### '):
                html_lines.append(f'<h3>{line[4:]}</h3>')
            elif line.startswith('**') and line.endswith('**'):
                html_lines.append(f'<p><strong>{line[2:-2]}</strong></p>')
            elif line.startswith('*') and line.endswith('*'):
                html_lines.append(f'<p><em>{line[1:-1]}</em></p>')
            elif line.startswith('- '):
                html_lines.append(f'<li>{line[2:]}</li>')
            elif line.startswith('---'):
                html_lines.append('<hr>')
            else:
                if line:
                    html_lines.append(f'<p>{line}</p>')
        
        return '\n'.join(html_lines)
    
    def _add_html_styling(self, html_content: str) -> str:
        """Add CSS styling to HTML"""
        css_style = """
        <style>
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f9f9f9;
        }
        
        .container {
            background-color: white;
            padding: 40px;
            border-radius: 10px;
            box-shadow: 0 0 20px rgba(0,0,0,0.1);
        }
        
        h1 {
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
            margin-bottom: 30px;
        }
        
        h2 {
            color: #34495e;
            border-left: 4px solid #3498db;
            padding-left: 15px;
            margin-top: 30px;
        }
        
        h3 {
            color: #34495e;
            margin-top: 25px;
        }
        
        hr {
            border: none;
            height: 2px;
            background: linear-gradient(to right, #3498db, #transparent);
            margin: 30px 0;
        }
        
        strong {
            color: #2c3e50;
        }
        
        em {
            color: #7f8c8d;
        }
        
        li {
            margin-bottom: 5px;
        }
        
        .footer {
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #ecf0f1;
            font-size: 0.9em;
            color: #7f8c8d;
        }
        </style>
        """
        
        return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Project Proposal</title>
    {css_style}
</head>
<body>
    <div class="container">
        {html_content}
    </div>
</body>
</html>"""
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported export formats"""
        return self.supported_formats.copy()
    
    def validate_format(self, format_name: str) -> bool:
        """Validate if format is supported"""
        return format_name.lower() in self.supported_formats